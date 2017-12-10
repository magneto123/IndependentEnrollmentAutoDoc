#-*- coding:utf-8 -*-
import os  
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH#段落格式之对齐方式
#from docx.enum.table import WD_ROW_HEIGHT_RULE#表格高度

import time

import readstuandunivinfo as rsui
import dataprocess


#设定行文字方向
def set_vert_cell_direction(cell):
    # https://github.com/python-openxml/python-docx/issues/55
    tc = cell._tc
    tcPr = tc.tcPr
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), 'btLr')
    tcPr.append(textDirection)

#设置某一行的高度
def set_row_height(row,heightv):
    # https://stackoverflow.com/questions/37532283/python-how-to-adjust-row-height-of-table-in-docx
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(heightv))
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)
    
    
#设置段落内容和格式
def SetParagraphTextAndFormat(paragraph,txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag):
    '''
    功能：设置表格内容和格式
    
    '''
    run = paragraph.add_run(txt)
    run.font.size = Pt(fontsize)#设置字号
    run.font.name = fontname#字体
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    run.font.italic = itaflag#是否斜体
    run.font.bold = boldflag#是否粗体
    #段落格式
    #左缩进
    paragraph.paragraph_format.left_indent = Pt(leftindent)#设置段落格式左缩进，单位为磅
    #首行缩进
    paragraph.paragraph_format.first_line_indent = Pt(firstlfind)#设置首行缩进
    
    if alignmentflag == -1:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT#设置段落格式为左对齐
    if alignmentflag == 0:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#设置段落格式为居中
    if alignmentflag == 1:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT#设置段落格式为右对齐
        
#设置表格内容和格式设置
def SetTableTextAndFormat(cell,txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag):
    '''
    功能：设置表格内容和格式
    
    '''
    paragraph = cell.paragraphs[0]#text = 
    SetParagraphTextAndFormat(paragraph,txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    
#写封面
def WriteFrontCoverPage(document,StuInfoData):
    '''
    功能：写报告的封面
    '''
    #先增加图像（公司logo图片为logo.png）
    document.add_picture(r'pictures\logo.png', height=Inches(1))
    #先写五行空行
    for i in range(2):
        document.add_paragraph('')
    
    #居中写产品名
    txt = u'北京展梦学业规划指导中心“一元咨询”产品'
    paragraph = document.add_paragraph()
    run = paragraph.add_run(txt)
    run.font.size = Pt(16)#设置字号
    run.font.name=u'黑体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#设置段落格式为居中
    #居中写报告名
    txt = u'自主招生院校专业定位及申报建议报告书'
    paragraph = document.add_paragraph()
    run = paragraph.add_run(txt)
    run.font.size = Pt(22)#设置字号
    run.font.name=u'黑体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#设置段落格式为居中
    
    #下面再输入1行空格
    for i in range(1):
        document.add_paragraph('')
        
    #下面居中输入报告的内容简介
    txt = u'本报告分为如下几个部分：\n（1）专业取向分析及建议\n（2）院校定位及建议\n（3）附件：自主招生内参资料'
    paragraph = document.add_paragraph()
    run = paragraph.add_run(txt)
    run.font.size = Pt(12)#设置字号
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    paragraph.paragraph_format.left_indent = Inches(2)#设置段落格式左缩进

    #下面再输入3行空格
    for i in range(2):
        document.add_paragraph('')
    
    
    #输入学生信息
    curtime = time.strftime('%Y-%m-%d',time.localtime(time.time()))
    name = StuInfoData['name'].decode('utf-8')
    school = StuInfoData['school'].decode('utf-8')
    grade = StuInfoData['grade'].decode('utf-8')
    txt = u'姓  名：' + name + u'\n学  校：' + school + u'\n年  级：' + grade + u'\n时  间：' + curtime
    paragraph = document.add_paragraph()
    run = paragraph.add_run(txt)
    run.font.size = Pt(14)#设置字号
    run.font.name=u'黑体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    paragraph.paragraph_format.left_indent = Inches(2)#设置段落格式左缩进
        
    #下面再输入3行空格
    for i in range(2):
        document.add_paragraph('')
        
    #输入备注信息：
    txt = u'本报告由北京展梦学业规划指导中心提供，仅对本次采集的数据负责，如有问题请致电400-88888888' + \
    u'或发送邮件至学业邮箱：xf700@qq.com 联系我们！'
    paragraph = document.add_paragraph()
    run = paragraph.add_run(txt)
    run.font.size = Pt(12)#设置字号
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    paragraph.paragraph_format.first_line_indent = Inches(0.3)
    
    #空行
    document.add_paragraph('')
    #祝福语
    txt = u'北京展梦学业规划指导中心祝莘莘学子心想事成，金榜题名！'
    paragraph = document.add_paragraph()
    run = paragraph.add_run(txt)
    run.font.size = Pt(12)#设置字号
    run.font.name=u'楷体'
    run.font.italic = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#设置段落格式为居中
    
    #分页
    document.add_page_break()
    
    return document

#写报告正文第一页，包括个人信息和专业建议结果
def WriteFirstPage(document,StuInfoData,MRTIData,SubMajData):
    '''
    功能：写报告正文第一页，包括个人信息和专业建议结果
    '''
    #写在前面的话：
    if StuInfoData['identity'].find('学生') > -1:
        txt = StuInfoData['name'].decode('utf-8') + u'同学，您好！由衷的感谢您选择北京展梦学业规划指导中心来为您规划学业，'
    else:
        txt = u'尊敬的家长，您好！由衷的感谢您选择北京展梦学业规划指导中心来为您的孩子规划学业，'
    txt += u'根据您提供的信息，我们为您的自主招生咨询作出如下建议，敬请参考，谢谢！'
    paragraph = document.add_paragraph()
    run = paragraph.add_run(txt)
    run.font.size = Pt(12)#设置字号
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    paragraph.paragraph_format.first_line_indent = Inches(0.3)
    
    #一级标题，个人信息及专业建议
    paragraph = document.add_heading(u'',1)
    SetParagraphTextAndFormat(paragraph,u'一、个人信息及专业建议',20,u'黑体',-1,0,0,False,False)
    #居中写报告名
    txt = u'表1 学生个人信息及专业建议'
    paragraph = document.add_paragraph()
    SetParagraphTextAndFormat(paragraph,txt,12,u'黑体',0,0,0,False,False)
    
    #增加表格写个人信息
    table = document.add_table(rows=10, cols=7,style = 'Table Grid')
    #设置表格部分行高度
    set_row_height(table.rows[6],3000)
    set_row_height(table.rows[7],1000)
    set_row_height(table.rows[8],1500)
    set_row_height(table.rows[9],1500)
    #设置表格部分列的宽度
    table.cell(0,0).width = Inches(0.3)
    table.cell(0,6).width = Inches(2)
    #合并单元格
    cell = table.cell(0,0)
    cell.merge(table.cell(4,0))#合并个人信息
    cell = table.cell(5,0)
    cell.merge(table.cell(8,0))#合并MBTI性格测试项    
    
    cell = table.cell(1,2)
    cell.merge(table.cell(1,4))#合并邮箱内容项
    cell = table.cell(3,2)
    cell.merge(table.cell(3,4))#合并优势学科内容项
    cell = table.cell(4,2)
    cell.merge(table.cell(4,4))#合并意向城市内容项
    cell = table.cell(5,2)
    cell.merge(table.cell(5,6))#合并性格类型内容项
    cell = table.cell(6,2)
    cell.merge(table.cell(6,6))#合并性格描述内容项
    cell = table.cell(7,2)
    cell.merge(table.cell(7,6))#合并适合领域内容项
    cell = table.cell(8,2)
    cell.merge(table.cell(8,6))#合并适合职业内容项
    cell = table.cell(9,1)
    cell.merge(table.cell(9,6))#合并建议专业内容项
    
    #表格字体
    fontsize = 10
    fontname = u'宋体'
    alignmentflag = 0
    leftindent = 0
    firstlfind = 0
    itaflag = False 
    boldflag = False
    #表格结构内容
    txt = u'\n\n个人\n信息'#个人信息
    SetTableTextAndFormat(table.cell(0,0),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'姓名'
    SetTableTextAndFormat(table.cell(0,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'性别'
    SetTableTextAndFormat(table.cell(0,3),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'手机'
    SetTableTextAndFormat(table.cell(0,5),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'邮箱'
    SetTableTextAndFormat(table.cell(1,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'学校'
    SetTableTextAndFormat(table.cell(1,5),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'分科'
    SetTableTextAndFormat(table.cell(2,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'年级'
    SetTableTextAndFormat(table.cell(2,3),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'成绩'
    SetTableTextAndFormat(table.cell(2,5),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'优势学科'
    SetTableTextAndFormat(table.cell(3,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'排斥学科'
    SetTableTextAndFormat(table.cell(3,5),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'意向城市'
    SetTableTextAndFormat(table.cell(4,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'意向高校水平'
    SetTableTextAndFormat(table.cell(4,5),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'\n\n\n\n\n\n\n\n\nMBTI\n性格\n测试'
    SetTableTextAndFormat(table.cell(5,0),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'性格类型'
    SetTableTextAndFormat(table.cell(5,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'\n\n\n\n\n性格描述'
    SetTableTextAndFormat(table.cell(6,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'\n\n适合领域'
    SetTableTextAndFormat(table.cell(7,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'\n\n适合职业'
    SetTableTextAndFormat(table.cell(8,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'\n建议\n选择\n的专\n业'
    SetTableTextAndFormat(table.cell(9,0),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    
    #表格内容
    txt = StuInfoData['name'].decode('utf-8')
    SetTableTextAndFormat(table.cell(0,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['gender'].decode('utf-8')
    SetTableTextAndFormat(table.cell(0,4),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['phonenum'].decode('utf-8')
    SetTableTextAndFormat(table.cell(0,6),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['email'].decode('utf-8')
    SetTableTextAndFormat(table.cell(1,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['school'].decode('utf-8')
    SetTableTextAndFormat(table.cell(1,6),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['artsorsci'].decode('utf-8')
    SetTableTextAndFormat(table.cell(2,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['grade'].decode('utf-8')
    SetTableTextAndFormat(table.cell(2,4),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['graderank'].decode('utf-8')
    SetTableTextAndFormat(table.cell(2,6),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['advantagesubject'].decode('utf-8')
    SetTableTextAndFormat(table.cell(3,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['disadvtsubject'].decode('utf-8')
    SetTableTextAndFormat(table.cell(3,6),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['intentunivlocation'].decode('utf-8')
    SetTableTextAndFormat(table.cell(4,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = StuInfoData['intentuniverrank'].decode('utf-8')
    SetTableTextAndFormat(table.cell(4,6),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)

    txt = MRTIData['type'].decode('utf-8')
    SetTableTextAndFormat(table.cell(5,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    #格式
    alignmentflag = -1
    firstlfind = 0#首行缩进已在原始数据中设置，所以不再设置

    txt = MRTIData['description'].decode('utf-8')
    SetTableTextAndFormat(table.cell(6,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = MRTIData['field'].decode('utf-8')
    SetTableTextAndFormat(table.cell(7,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = MRTIData['profession'].decode('utf-8')
    SetTableTextAndFormat(table.cell(8,2),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    
    #写入专业数据
    txt = SubMajData.decode('utf-8')
    SetTableTextAndFormat(table.cell(9,1),txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    #分页
    document.add_page_break()
    return document


    
#写学校信息
def WriteUniversityInfoPage(document,StuUniversityApplyConditionsData):
    '''
    功能：写学校信息
    '''
    #一级标题，推荐的学校
    paragraph = document.add_heading(u'',1)
    SetParagraphTextAndFormat(paragraph,u'二、自主招生院校定位及申报建议',20,u'黑体',-1,0,0,False,False)
    #段前说明
    txt = u'根据您提供的数据，我们推荐您考虑以下院校的自主招生。需要请您注意的是2018年度自招的信息在2018年3月发布，'+\
    u'为方便您备考，下面给出相应学校在2017年度的自招信息。'
    paragraph = document.add_paragraph()
    SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,False)#首行缩进20磅
    #居中写表名
    txt = u'表2 自主招生院校定位'
    paragraph = document.add_paragraph()
    SetParagraphTextAndFormat(paragraph,txt,12,u'黑体',0,0,0,False,False)
    #增加表格写学校信息
    #先设置表格字体
    fontsize = 10
    fontname = u'宋体'
    alignmentflag = 0
    leftindent = 0
    firstlfind = 0
    itaflag = False 
    boldflag = False
    #创建表格
    table = document.add_table(rows=1, cols=6,style = 'Table Grid')
    hdr_cells = table.rows[0].cells
    
    txt = u'序号'
    SetTableTextAndFormat(hdr_cells[0],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'学校名称'
    SetTableTextAndFormat(hdr_cells[1],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'地点'
    SetTableTextAndFormat(hdr_cells[2],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'学校层次'
    SetTableTextAndFormat(hdr_cells[3],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'限报学校数'
    SetTableTextAndFormat(hdr_cells[4],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    txt = u'招生人数'
    SetTableTextAndFormat(hdr_cells[5],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    #再增加表格元素
    schoolnum = len(StuUniversityApplyConditionsData)
    for i in xrange(schoolnum):
        row_cells = table.add_row().cells
        txt = str(i+1).decode('utf-8')
        SetTableTextAndFormat(row_cells[0],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
        txt = StuUniversityApplyConditionsData[i]['schoolname'].decode('utf-8')
        SetTableTextAndFormat(row_cells[1],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
        txt = StuUniversityApplyConditionsData[i]['schoollocation'].decode('utf-8')
        SetTableTextAndFormat(row_cells[2],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
        txt = StuUniversityApplyConditionsData[i]['schoolrank'].decode('utf-8')
        SetTableTextAndFormat(row_cells[3],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
        txt = StuUniversityApplyConditionsData[i]['schoollimit'].decode('utf-8')
        SetTableTextAndFormat(row_cells[4],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
        txt = StuUniversityApplyConditionsData[i]['plan'].decode('utf-8')
        SetTableTextAndFormat(row_cells[5],txt,fontsize,fontname,alignmentflag,leftindent,firstlfind,itaflag,boldflag)
    
    #分段写每个学校的招考要求
    txt = u'上表是北京展梦学业规划指导中心给您的自主招生院校定位，下面详细给出每所学校的招考要求，敬请您参考，谢谢！'
    paragraph = document.add_paragraph()
    SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,False)#首行缩进20磅
    ##循环填写学校的其他信息
    for i in range(schoolnum):
        ##二级标题，学校名称
        txt = (str(i+1) + '. ' + StuUniversityApplyConditionsData[i]['schoolname']).decode('utf-8')
        paragraph = document.add_heading(u'',2)
        SetParagraphTextAndFormat(paragraph,txt,16,u'黑体',-1,0,0,False,False)
        #分条目写其他信息
        infoi = 1#报名时间
        paragraph = document.add_paragraph()
        txt = u'（' + str(infoi) + u'）报名时间：'
        SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,True)
        txt = StuUniversityApplyConditionsData[i]['applytime'].decode('utf-8')
        SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,False)
        #邮寄材料截止时间
        infoi += 1
        paragraph = document.add_paragraph()
        txt = u'（' + str(infoi) + u'）邮寄材料截止时间：'
        SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,True)
        txt = StuUniversityApplyConditionsData[i]['materialpostdeadline'].decode('utf-8')
        SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,False)
        #笔面试时间
        infoi += 1
        paragraph = document.add_paragraph()
        txt = u'（' + str(infoi) + u'）笔面试时间：'
        SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,True)
        txt = StuUniversityApplyConditionsData[i]['examtime'].decode('utf-8')
        SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,False)
        #专利论文要求：
        paper = StuUniversityApplyConditionsData[i]['paper'].decode('utf-8')
        patent = StuUniversityApplyConditionsData[i]['patent'].decode('utf-8')
        if len(paper) > 0 or len(patent) > 0:
            infoi += 1
            paragraph = document.add_paragraph()
            txt = u'（' + str(infoi) + u'）专利论文要求：'
            SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,True)
            
            infoix = 0
            if len(paper) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 论文：' + paper
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(patent) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 专利：' + patent
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
        #学科竞赛要求
        #先判断有无该项要求
        math = StuUniversityApplyConditionsData[i]['math'].decode('utf-8')
        phis = StuUniversityApplyConditionsData[i]['physics'].decode('utf-8')
        chem = StuUniversityApplyConditionsData[i]['chemistry'].decode('utf-8')
        bio = StuUniversityApplyConditionsData[i]['biology'].decode('utf-8')
        info = StuUniversityApplyConditionsData[i]['infomatics'].decode('utf-8')
        if len(math) > 0 or len(phis) > 0 or len(chem) > 0 or len(bio) > 0 or len(info) > 0:
            infoi += 1
            paragraph = document.add_paragraph()
            txt = u'（' + str(infoi) + u'）学科竞赛要求：'
            SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,True)
            
            infoix = 0
            if len(math) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 数学：' + math
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(phis) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 物理：' + phis
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(chem) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 化学：' + chem
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(bio) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 生物：' + bio
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(info) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 信息学：' + info
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
                
        #文科类竞赛要求
        newcon = StuUniversityApplyConditionsData[i]['newconceptcomposition'].decode('utf-8')
        innocom = StuUniversityApplyConditionsData[i]['innovativecomposition'].decode('utf-8')
        chinesenews = StuUniversityApplyConditionsData[i]['chinesenewspapercup'].decode('utf-8')
        yeshengtao = StuUniversityApplyConditionsData[i]['yeshengtaocup'].decode('utf-8')
        peking = StuUniversityApplyConditionsData[i]['pekinguculture'].decode('utf-8')
        innoeng = StuUniversityApplyConditionsData[i]['innovativeenglish'].decode('utf-8')
        engcom = StuUniversityApplyConditionsData[i]['englishcompetition'].decode('utf-8')
        if len(newcon) > 0 or len(innocom) > 0 or len(chinesenews) > 0 or len(yeshengtao) > 0 or len(peking) > 0 \
            or len(innoeng) > 0 or len(engcom) > 0:
            infoi += 1
            paragraph = document.add_paragraph()
            txt = u'（' + str(infoi) + u'）文科类竞赛要求：'
            SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,True)
            
            infoix = 0
            if len(newcon) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 新概念作文大赛：' + newcon
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(innocom) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 创新作文大赛：' + innocom
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(chinesenews) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 语文报杯：' + chinesenews
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(yeshengtao) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 叶圣陶杯：' + yeshengtao
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(peking) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 北大学培文化：' + peking
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(innoeng) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 创新英语竞赛：' + innoeng
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(engcom) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 英语能力竞赛：' + engcom
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
        
        #科技创新活动大赛
        young = StuUniversityApplyConditionsData[i]['youngsciinnocom'].decode('utf-8')
        tomm = StuUniversityApplyConditionsData[i]['tomorrowscientist'].decode('utf-8')
        prim = StuUniversityApplyConditionsData[i]['primsecdschcomputer'].decode('utf-8')
        intsci = StuUniversityApplyConditionsData[i]['intenationalscieng'].decode('utf-8')
        intenv = StuUniversityApplyConditionsData[i]['intenationalenvironmentsciprj'].decode('utf-8')

        if len(young) > 0 or len(tomm) > 0 or len(prim) > 0 or len(intsci) > 0 or len(intenv) > 0:
            infoi += 1
            paragraph = document.add_paragraph()
            txt = u'（' + str(infoi) + u'）科技创新竞赛要求：'
            SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,True)
            
            infoix = 0
            if len(young) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 青少年科技创新大赛：' + young
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(tomm) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 明天小小科学技术科创大赛：' + tomm
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(prim) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 中小学电脑制作大赛：' + prim
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(intsci) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 国际科学与工程竞赛：' + intsci
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
            if len(intenv) > 0:
                paragraph = document.add_paragraph()
                infoix += 1
                txt = u'[' + str(infoix) + u'] 国际环境科研项目：' + intenv
                SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,30,20,False,False)
                
        ###############################################
    #写结束语
    #下面再输入3行空行
    for i in range(3):
        document.add_paragraph('')
    #段前说明
    txt = u'上述内容为北京展梦学业规划指导中心根据您提供的信息为您提供的自主招生学校定位信息，敬请参考！如对结果有疑问，请致电400-88888888' + \
    u'或发送邮件至学业邮箱：xf700@qq.com 联系北京展梦学业规划指导中心的黄老师（学业规划）或侯老师（自主招生）！'
    paragraph = document.add_paragraph()
    SetParagraphTextAndFormat(paragraph,txt,12,u'宋体',-1,0,20,False,False)#首行缩进20磅    
    #学校信息页结束，加分页
    #document.add_page_break()
    return document
    
    
#将信息写入word文件总调用函数
def WriteStuAndUinversityDatatoDoc(FilePathStu,iStu):
    '''
    功能：写doc文件总调用函数
    '''
    curospath = os.path.abspath('.')
    FilePathMBTI = curospath + '//data//MBTI_Data.xlsx'

    flag,MBTIData = rsui.ReadMBTIDataFromExcel(FilePathMBTI)
    if flag == 0:
        return 0
    #读取不同学生和学校的报考信息
    StuData,sugmajor,stumbti,stuanduniverinfo = dataprocess.GetAllStudentsUniversityandApplyConditions(FilePathStu,iStu)        
    
    stunum = len(StuData)
    
    LoopList = []
    if len(iStu) == 1:
        if iStu[0] >= 0 and iStu[0] < stunum:
            LoopList.append(iStu[0])
        else:
            LoopList = range(stunum)
    else:
        #检查有效性
        for i in range(len(iStu)):
            if iStu[i] < 0 or iStu[i] >= stunum:
                del iStu[i]
        LoopList = iStu
        
    ip = 0
    for i in LoopList:#循环写入
        document = Document()
        WriteFrontCoverPage(document,StuData[i])#写封面页
        WriteFirstPage(document,StuData[i],MBTIData[stumbti[ip]],sugmajor[ip])#写第一页
        WriteUniversityInfoPage(document,stuanduniverinfo[ip])#写学校的信息
        ip += 1
        docpath = curospath + '\\报告\\' + StuData[i]['name'] + '.docx'#str(i) + '.docx'#
        docpath = docpath.decode('utf-8')
        document.save(docpath)
    return 1

'''
if __name__=="__main__":
    FilePathStu = r'C:\Users\CXJ-PC\Desktop\data\stuinfo.xlsx'
    WriteStuAndUinversityDatatoDoc(FilePathStu,[-1])
'''